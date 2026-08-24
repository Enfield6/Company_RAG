from io import BytesIO
from unittest.mock import PropertyMock, patch

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from PIL import Image

from app.documents.docx_parser import DocxParser


def test_docx_parser_preserves_image_neighborhood(tmp_path) -> None:
    image_buffer = BytesIO()
    Image.new("RGB", (40, 20), color=(16, 163, 127)).save(image_buffer, format="PNG")
    image_buffer.seek(0)

    document = Document()
    document.add_heading("部署流程", level=1)
    document.add_paragraph("先准备服务器。")
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(image_buffer, width=Inches(1))
    document.add_paragraph("然后执行安装并检查健康状态。")
    path = tmp_path / "guide.docx"
    document.save(path)

    elements = DocxParser().parse(path)
    image = next(element for element in elements if element.kind == "image")

    assert image.image_bytes
    assert image.heading_path == ["部署流程"]
    assert image.metadata["previous_text"] == "先准备服务器。"
    assert image.metadata["next_text"] == "然后执行安装并检查健康状态。"
    assert elements[image.sequence_no] is image


def test_docx_parser_accepts_paragraph_without_resolved_style(tmp_path) -> None:
    document = Document()
    document.add_paragraph("这个段落的样式定义无法解析。")
    path = tmp_path / "missing-style.docx"
    document.save(path)

    with patch.object(Paragraph, "style", new_callable=PropertyMock, return_value=None):
        elements = DocxParser().parse(path)

    assert len(elements) == 1
    assert elements[0].text == "这个段落的样式定义无法解析。"
